"""Generate DOCX and PDF from final markdown content.

Resume: python-docx — inject tailored summary + skills into the template by
        anchoring on section headings → 1-page check → LibreOffice PDF
Cover letter: python-docx body replacement → LibreOffice PDF
"""
import copy
import re
import shutil
import subprocess
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph


def _apply_calibri(doc: Document) -> None:
    """Force Calibri on every run in the document to ensure font consistency."""
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "Calibri"
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                rFonts.set(qn(attr), "Calibri")


_SOFFICE_CANDIDATES = [
    "soffice",
    "libreoffice",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/usr/bin/libreoffice",
    "/usr/local/bin/libreoffice",
]


def _find_soffice() -> str:
    for candidate in _SOFFICE_CANDIDATES:
        if shutil.which(candidate) or Path(candidate).exists():
            return candidate
    raise RuntimeError(
        "LibreOffice not found. Install it with: brew install --cask libreoffice\n"
        "Then restart the backend server."
    )


def _libreoffice_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    soffice = _find_soffice()
    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", str(docx_path), "--outdir", str(out_dir)],
        capture_output=True, text=True, timeout=120,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice failed: {result.stderr}")
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise RuntimeError(f"PDF not found at {pdf_path}")
    return pdf_path


def _page_count(pdf_path: Path) -> int:
    result = subprocess.run(
        ["pdfinfo", str(pdf_path)], capture_output=True, text=True, timeout=30,
    )
    match = re.search(r"Pages:\s+(\d+)", result.stdout)
    return int(match.group(1)) if match else -1


_GREETING_PATTERNS = [
    r"sehr geehrte",
    r"dear ",
]

_CLOSING_PATTERNS = [
    r"mit freundlichen grüßen",
    r"kind regards",
    r"freundliche grüße",
    r"best regards",
    r"yours sincerely",
    r"hochachtungsvoll",
]

_CONTACT_PATTERN = re.compile(r"[\w.+-]+@[\w-]+\.\w+|[\+\d][\d\s\(\)\-\.]{6,}")


def _strip_letter_boilerplate(text: str) -> str:
    """Remove greeting, truncate at the closing salutation, but preserve any
    contact/availability line that appears after the closing (LLMs sometimes
    place it there)."""
    lines = text.split("\n")
    body_lines: list[str] = []
    tail_lines: list[str] = []  # lines after the closing
    hit_closing = False

    for line in lines:
        lower = line.strip().lower()
        if any(re.match(p, lower) for p in _GREETING_PATTERNS):
            continue
        if any(re.match(p, lower) for p in _CLOSING_PATTERNS):
            hit_closing = True
            continue  # skip the closing line itself
        if hit_closing:
            tail_lines.append(line)
        else:
            body_lines.append(line)

    # Rescue any line after the closing that contains an email or phone
    rescued = [l for l in tail_lines if _CONTACT_PATTERN.search(l)]
    if rescued:
        body_lines.extend([""] + rescued)

    return "\n".join(body_lines)


# Section headings used to anchor summary/skills replacement in the resume
# template DOCX. The templates carry no paraId/style markers (every paragraph
# uses the same "Body A" style), so we locate sections by their visible heading
# text. Keep these in sync with the template DOCX headings per language.
RESUME_HEADINGS = {
    "en": {"summary": "PROFESSIONAL SUMMARY", "skills": "TECHNICAL SKILLS", "after_skills": "PROFESSIONAL EXPERIENCE"},
    "de": {"summary": "PROFIL", "skills": "TECHNISCHE KENNTNISSE", "after_skills": "BERUFSERFAHRUNG"},
}

def _extract_sections(resume_md: str) -> dict[str, str]:
    """Pull the tailored profile summary and raw skills markdown from the resume."""
    # Terminate each section at the next TOP-LEVEL heading ("# "), not at the
    # "## Group" sub-headings inside the skills section.
    summary_match = re.search(r"# Profile\n(.*?)(?=\n#(?!#)|\Z)", resume_md, re.DOTALL)
    skills_match = re.search(r"# Skills\n(.*?)(?=\n#(?!#)|\Z)", resume_md, re.DOTALL)
    return {
        "summary": (summary_match.group(1).strip() if summary_match else ""),
        "skills_md": (skills_match.group(1).strip() if skills_match else ""),
    }


def _parse_skills_groups(skills_md: str) -> list[tuple[str, str]]:
    """Parse the markdown skills section into (label, items) pairs.

    Handles the master '## Group\\nitems' format and degrades gracefully to
    freeform 'Label: items' lines (or label-less lines) when the LLM deviates.
    """
    text = skills_md.strip()
    if not text:
        return []
    groups: list[tuple[str, str]] = []
    for block in re.split(r"\n(?=##\s)", text):
        m = re.match(r"##\s+(.+?)\n(.+)", block.strip(), re.DOTALL)
        if m:
            label = m.group(1).strip()
            items = " ".join(ln.strip() for ln in m.group(2).splitlines() if ln.strip())
            groups.append((label, items))
    if groups:
        return groups
    # Fallback: no '##' headers — treat each non-empty line as its own entry
    for line in text.splitlines():
        line = line.strip().lstrip("-*").strip()
        if not line:
            continue
        if ":" in line:
            label, _, items = line.partition(":")
            groups.append((label.strip(), items.strip()))
        else:
            groups.append(("", line))
    return groups


def _set_paragraph_text(para: Paragraph, text: str) -> None:
    """Put `text` into the paragraph's first run and blank the rest, preserving
    the run's formatting (font/colour/size)."""
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for run in para.runs[1:]:
        run.text = ""


def _set_skill_runs(p_element, label: str, items: str) -> None:
    """On a cloned skill paragraph element, set the bold-label run and the list
    run. The template skill paragraph has two <w:t> runs (label + items)."""
    t_elems = p_element.findall(".//" + qn("w:t"))
    values = [f"{label}: " if label else "", items]
    for i, t in enumerate(t_elems):
        t.text = values[i] if i < len(values) else ""
        t.set(qn("xml:space"), "preserve")


def _recolour_docx(path: Path, old: str, new: str) -> None:
    """Wholesale-replace an accent colour token in word/document.xml in place."""
    import zipfile

    tmp = path.with_suffix(".recolour.docx")
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                data = data.replace(old.encode(), new.encode())
            zout.writestr(item, data)
    tmp.replace(path)


def generate_resume_docx(
    resume_md: str,
    template_path: Path,
    output_path: Path,
    language: str = "en",
    accent_colour: str | None = None,
) -> Path:
    sections = _extract_sections(resume_md)
    headings = RESUME_HEADINGS.get(language, RESUME_HEADINGS["en"])
    groups = _parse_skills_groups(sections["skills_md"])

    doc = Document(str(template_path))
    paras = doc.paragraphs

    def _find(heading: str) -> int | None:
        target = heading.strip().lower()
        for i, p in enumerate(paras):
            if p.text.strip().lower() == target:
                return i
        return None

    # ── Summary: the single paragraph after the summary heading ───────────────
    sum_i = _find(headings["summary"])
    if sum_i is not None and sum_i + 1 < len(paras) and sections["summary"]:
        _set_paragraph_text(paras[sum_i + 1], sections["summary"])

    # ── Skills: rebuild the paragraphs between skills and the next heading ─────
    skills_i = _find(headings["skills"])
    exp_i = _find(headings["after_skills"])
    if skills_i is not None and exp_i is not None and exp_i > skills_i + 1 and groups:
        skill_paras = paras[skills_i + 1:exp_i]
        # Clone the first existing skill paragraph per group to keep run styling
        # (bold label run + regular list run), inserting after the last original.
        tmpl_p = skill_paras[0]._p
        prev = skill_paras[-1]._p
        for label, items in groups:
            new_p = copy.deepcopy(tmpl_p)
            _set_skill_runs(new_p, label, items)
            prev.addnext(new_p)
            prev = new_p
        for p in skill_paras:
            p._p.getparent().remove(p._p)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    if accent_colour:
        _recolour_docx(output_path, "1a56a4", accent_colour)

    return output_path


def generate_cover_letter_docx(
    cover_letter_md: str,
    template_path: Path,
    output_path: Path,
    job_title: str,
    company: str,
    company_address: str,
    language: str,
) -> Path:
    doc = Document(str(template_path))
    date_str = datetime.now().strftime("%d.%m.%Y")

    if language == "en":
        subject = f"Application for {job_title}"
        greeting = "Dear Hiring Manager,"
        closing_text = "Kind regards,"
    else:
        subject = f"Bewerbung als {job_title}"
        greeting = "Sehr geehrte Damen und Herren,"
        closing_text = "Mit freundlichen Grüßen,"

    # Split address into lines (street on first, city on second)
    addr_lines = [ln.strip() for ln in company_address.splitlines() if ln.strip()]

    # ── First pass: replace fixed placeholders ────────────────────────────────
    paras_to_remove = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "[Unternehmensname]":
            para.text = company
        elif text == "[Straße Nr.]":
            if not addr_lines:
                paras_to_remove.append(para)
            else:
                # The last line is the city ([PLZ Stadt]); everything before it
                # is the street block. Any middle lines are inserted as extra
                # street paragraphs so a 3+ line address is not truncated.
                street_lines = addr_lines[:-1] if len(addr_lines) > 1 else addr_lines
                para.text = street_lines[0]
                anchor = para._p
                for extra in street_lines[1:]:
                    new_p = copy.deepcopy(para._p)
                    anchor.addnext(new_p)
                    anchor = new_p
                    Paragraph(new_p, para._parent).text = extra
        elif text == "[PLZ Stadt]":
            if len(addr_lines) > 1:
                para.text = addr_lines[-1]
            else:
                paras_to_remove.append(para)
        elif "Bewerbung als" in text or "Application as" in text or "Application for" in text:
            para.text = subject
        elif re.search(r'[^\d,]+,\s*\d{2}\.\d{2}\.\d{4}', text):
            # Keep the sender's place from the template, only refresh the date.
            # Translate München → Munich for English letters.
            m = re.search(r'([^\d,]+),\s*\d{2}\.\d{2}\.\d{4}', text)
            place = m.group(1).strip() if m else "München"
            if language == "en" and place.lower() in ("münchen", "muenchen"):
                place = "Munich"
            para.text = re.sub(r'[^\d,]+,\s*\d{2}\.\d{2}\.\d{4}', f"{place}, {date_str}", text)
        elif "Sehr geehrte" in text:
            para.text = greeting
        elif "Mit freundlichen Grüßen" in text and language == "en":
            para.text = closing_text

    for para in paras_to_remove:
        para._p.getparent().remove(para._p)

    # ── Second pass: replace body between greeting and closing ────────────────
    greeting_para = None
    closing_para = None
    old_body = []
    in_body = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if text == greeting:
            greeting_para = para
            in_body = True
            continue
        if "Mit freundlichen Grüßen" in text or "Kind regards" in text:
            closing_para = para
            in_body = False
            break
        if in_body:
            old_body.append(para)

    # Save body style before removing old paragraphs
    body_style = old_body[0].style if old_body else None

    # Remove old body paragraphs from XML
    for para in old_body:
        para._p.getparent().remove(para._p)

    # Insert new body paragraphs before the closing
    if closing_para:
        clean_body = _strip_letter_boilerplate(cover_letter_md)
        new_paras = [p.strip() for p in clean_body.strip().split("\n\n") if p.strip()]
        for p_text in new_paras:
            new_p = closing_para.insert_paragraph_before(p_text)
            if body_style:
                new_p.style = body_style
            new_p.paragraph_format.space_after = Pt(8)

    _apply_calibri(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _person_name_from_template(template: Path, language: str) -> str:
    """Strip the document-type suffix from the template stem to get the person name.

    e.g. 'FirstName_LastName_CV' -> 'FirstName_LastName'
    Falls back to the full stem if no known suffix is found.
    """
    stem = template.stem
    for suffix in ("_CV", "_Lebenslauf", "_Resume"):
        if stem.endswith(suffix):
            return stem[: -len(suffix)]
    return stem


def build_pdfs(
    resume_md: str,
    cover_letter_md: str,
    job_title: str,
    company: str,
    company_address: str,
    language: str,
    template_resume: Path,
    template_cover: Path,
    output_dir: Path,
    person_name: str = "",
) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)

    if not person_name:
        person_name = _person_name_from_template(template_resume, language)

    if language == "de":
        cv_name = f"{person_name}_Lebenslauf"
        cl_name = f"{person_name}_Anschreiben"
    else:
        cv_name = f"{person_name}_CV"
        cl_name = f"{person_name}_Cover_Letter"

    cv_docx = output_dir / f"{cv_name}.docx"
    cl_docx = output_dir / f"{cl_name}.docx"

    generate_resume_docx(resume_md, template_resume, cv_docx, language)
    generate_cover_letter_docx(cover_letter_md, template_cover, cl_docx, job_title, company, company_address, language)

    cv_pdf = _libreoffice_to_pdf(cv_docx, output_dir)
    pages = _page_count(cv_pdf)
    cv_page_warning = None
    if pages != 1:
        cv_page_warning = (
            f"CV is {pages} pages — must be exactly 1. "
            "Shorten the profile summary (< 200 chars) or trim skills/bullets."
        )

    cl_pdf = _libreoffice_to_pdf(cl_docx, output_dir)

    return {
        "resume_docx": str(cv_docx),
        "resume_pdf": str(cv_pdf),
        "cover_letter_docx": str(cl_docx),
        "cover_letter_pdf": str(cl_pdf),
        "cv_page_warning": cv_page_warning,
    }
